"""Resume generation service.

Собирает адаптированное под конкретную вакансию резюме:
1. Sonnet генерирует структурированный JSON с полями резюме
2. python-docx собирает готовый .docx по нашему шаблону-стилю

Стиль имитирует эталонное резюме Ульяны:
- Calibri, чёрно-серая палитра, one-column, ATS-friendly
- Заголовки разделов заглавными с тонкой линией снизу
- Опыт: жирный заголовок роли, курсивом компания+даты,
  короткое описание, буллеты с жирным началом
"""

from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from typing import Any

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches

from src.services.claude import ClaudeService
from src.sources.base import Vacancy
from datetime import datetime, timezone

log = structlog.get_logger(__name__)

RESUME_MODEL = "claude-sonnet-4-6"

MAX_VACANCY_DESCRIPTION_CHARS = 4000
MAX_CV_SUMMARY_CHARS = 3000


RESUME_SYSTEM_PROMPT = """\
Ты — эксперт по составлению резюме уровня senior HR-консультанта. \
Задача: собрать резюме конкретного человека под конкретную вакансию, \
которое пройдёт ATS-фильтры и вызовет желание пригласить на интервью.

СЕГОДНЯШНЯЯ ДАТА: {current_date}. При расчёте опыта работы \
(например "N лет опыта", "работает с YYYY года") используй эту дату, \
не свои представления о времени. Если в резюме указано "с ноября 2021 по н.в." \
и сегодня 28 июля 2026 — это 4 года 8 месяцев опыта, а не 3.5.

ЖЁСТКИЕ ПРАВИЛА:

1. ЯЗЫК: определи по языку описания вакансии. Русскоязычная вакансия — резюме на русском. \
Английская — на английском. Смешанная — на языке большинства текста.

2. ФАКТОЛОГИЯ: используй только факты из профиля и загруженных резюме. \
Не выдумывай компании, роли, цифры, даты. Если данных нет — обходи их, \
не заполняй placeholder'ами.

2a. КОНТАКТЫ: обязательно найди в тексте загруженных резюме (cv_sources) \
email, телефон, Telegram, LinkedIn, город/страну проживания. Вставь их в шапку резюме. \
Если в резюме контактов нет — оставь только имя и headline без пустых плейсхолдеров.

3. АДАПТАЦИЯ ПОД ВАКАНСИЮ:
- Проанализируй ключевые требования вакансии (обязательные навыки, что важно бизнесу)
- Расставь акценты на то, что максимально совпадает с этими требованиями
- В summary укажи целевую роль близко к формулировке из вакансии
- В буллетах используй ключевые слова из вакансии там, где они честно подходят

4. СТРУКТУРА (это порядок разделов в резюме):
- ШАПКА: имя, professional headline (одна строка про роль), контакты
- ПРОФИЛЬ: 3-4 предложения. Кто человек как специалист, ключевая ценность для работодателя, \
опыт в конкретных доменах. Без клише ("коммуникабельный", "командный игрок"). \
Без воды.
- ОПЫТ РАБОТЫ: от последнего к первому. Для каждой роли:
  * Название роли (как звучит в резюме, часто отличается от формального job title)
  * Компания, город (если есть), даты
  * 1-2 предложения описания роли — контекст и уровень ответственности
  * 3-5 буллетов достижений. Формат: "Жирный заголовок достижения: развёрнутое описание с цифрами и результатом"
- ОБРАЗОВАНИЕ: одна-две строки. Степень, вуз, годы. Только если есть.
- НАВЫКИ: 4-6 буллетов группами. Формат: "Категория: список навыков через запятую"

5. БУЛЛЕТЫ (критически важно):
- Каждый буллет — конкретное достижение, не обязанность
- Формат: "Сделал X → результат Y". Не "занимался Y".
- Цифры везде где есть в данных (% роста, объёмы, размеры команд, бюджеты). \
Не выдумывай цифры если их нет — просто пропусти их.
- Начинай с ГЛАГОЛА действия (разработал, внедрил, увеличил, построил, запустил)
- Первая часть буллета — короткая жирная фраза-заголовок (2-6 слов) с двоеточием, \
потом развёрнутое описание с деталями

6. СТИЛЬ:
- Активные глаголы, прошедшее время для прошлых ролей, настоящее для текущей
- Профессиональный но живой язык, не канцелярит
- Никаких "коммуникабельный", "стрессоустойчивый", "командный игрок", "ответственный"
- Никаких эмодзи, звёздочек, декоративных символов

7. ATS-ОПТИМИЗАЦИЯ:
- Стандартные заголовки разделов на языке резюме
- Ключевые слова из вакансии — там где честно подходят к опыту

ФОРМАТ ОТВЕТА:

Верни СТРОГО валидный JSON без обёрток (без ```json```) со следующей структурой:

{
  "language": "ru" или "en",
  "full_name": "Имя Фамилия",
  "professional_headline": "Роль / Специализация (короткая строка)",
  "contacts": "tg: @username · email@example.com (одной строкой то, что доступно)",
  "summary": "3-4 предложения профиля",
  "experience": [
    {
      "role_title": "Название роли",
      "company_line": "Компания, Город · Даты",
      "role_summary": "1-2 предложения контекста",
      "bullets": [
        {"bold": "Жирный заголовок: ", "text": "развёрнутое описание с результатом"}
      ]
    }
  ],
  "education": [
    {"degree": "Степень / программа", "institution": "Вуз · Годы"}
  ],
  "skills": [
    {"bold": "Категория: ", "text": "список навыков через запятую"}
  ],
  "diagnostics": {
    "match_percent": число от 0 до 100,
    "strong_alignment": ["короткая формулировка", "..."],
    "gaps": ["короткая формулировка", "..."]
  }
}

ДИАГНОСТИКА:
- match_percent: честная оценка того, насколько кандидат подходит вакансии
- strong_alignment: 2-3 фразы что сильно совпало (опыт, навыки, домены)
- gaps: 1-2 честных пробела (чего не хватает, что не покрыто). Может быть пустым если совпадение идеальное.

Верни ТОЛЬКО JSON, без единой строчки комментария до или после.
"""


@dataclass
class ResumeResult:
    """Результат генерации резюме."""
    docx_bytes: bytes
    match_percent: int
    strong_alignment: list[str]
    gaps: list[str]
    language: str  # "ru" / "en"


# --- Хелперы для промпта ---

def _extract_cv_summary(profile_data: dict, max_sources: int = 3) -> str:
    cv_sources = profile_data.get("cv_sources") or []
    if not isinstance(cv_sources, list):
        return ""
    summaries: list[str] = []
    for cv in cv_sources[:max_sources]:
        if not isinstance(cv, dict):
            continue
        summary = cv.get("summary_extracted") or ""
        if summary:
            summaries.append(str(summary))
    combined = "\n\n---\n\n".join(summaries)
    return combined[:MAX_CV_SUMMARY_CHARS]


def _build_user_message(vacancy: Vacancy, profile_data: dict) -> str:
    description = (vacancy.description or "")[:MAX_VACANCY_DESCRIPTION_CHARS]
    cv_summary = _extract_cv_summary(profile_data)

    profile_bits: list[str] = []
    for field in (
        "expertise", "target_roles", "current_role_summary",
        "seniority", "industries_interested", "location_preferences",
        "compensation", "ideal_work_description", "interests_and_resonance",
        "languages", "must_haves",
    ):
        val = profile_data.get(field)
        if not val:
            continue
        if isinstance(val, list):
            val = ", ".join(str(x) for x in val if x)
        elif isinstance(val, dict):
            val = ", ".join(f"{k}={v}" for k, v in val.items() if v)
        profile_bits.append(f"{field}: {val}")
    profile_block = "\n".join(profile_bits) if profile_bits else "(профиль пустой)"

    return f"""ПРОФИЛЬ ЧЕЛОВЕКА:
{profile_block}

РЕЗЮМЕ (краткое содержание загруженных резюме):
{cv_summary if cv_summary else "(резюме не загружено — опирайся только на профиль)"}

ВАКАНСИЯ, ПОД КОТОРУЮ АДАПТИРУЕМ:
Название: {vacancy.title}
Компания: {vacancy.company or "не указана"}
Локация: {vacancy.location or "не указана"}
Зарплата: {vacancy.salary or "не указана"}

Описание:
{description}

Собери резюме под эту вакансию и верни JSON по схеме."""


# --- Хелперы для DOCX ---

FONT_NAME = "Calibri"
COLOR_HEADING = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)
COLOR_LINE = "999999"


def _set_run_font(run, size_pt: float, bold: bool = False, italic: bool = False,
                  color: RGBColor | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # Явно прописываем шрифт для CS/EA/Complex script (кириллица считается EA в некоторых системах)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT_NAME)


def _add_paragraph_border_bottom(paragraph) -> None:
    """Тонкая линия снизу параграфа (для заголовков секций)."""
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLOR_LINE)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_header(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    _set_run_font(run, 11, bold=True, color=COLOR_HEADING)
    _add_paragraph_border_bottom(p)


def _add_simple_paragraph(doc, text: str, size_pt: float = 10,
                          bold: bool = False, italic: bool = False,
                          color: RGBColor | None = None,
                          space_after: float = 5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_run_font(run, size_pt, bold=bold, italic=italic, color=color or COLOR_BODY)


def _add_bullet(doc, bold_part: str, rest_text: str) -> None:
    """Буллет с жирным началом и обычным продолжением."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run_bold = p.add_run(bold_part)
    _set_run_font(run_bold, 10, bold=True, color=COLOR_HEADING)
    run_rest = p.add_run(rest_text)
    _set_run_font(run_rest, 10, color=COLOR_BODY)


def _build_docx(resume_data: dict) -> bytes:
    """Собирает .docx по нашей структуре из JSON от Sonnet."""
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.625)
        section.right_margin = Inches(0.625)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(10)

    # --- ШАПКА ---
    name = resume_data.get("full_name", "")
    headline = resume_data.get("professional_headline", "")
    contacts = resume_data.get("contacts", "")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    _set_run_font(run, 16, bold=True, color=COLOR_HEADING)

    if headline:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        run = p.add_run(headline)
        _set_run_font(run, 11, color=COLOR_MUTED)

    if contacts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(contacts)
        _set_run_font(run, 10, color=COLOR_MUTED)

    # Заголовок секции по языку
    lang = resume_data.get("language", "ru")
    section_labels = {
        "ru": {"profile": "Профиль", "experience": "Опыт работы",
               "education": "Образование", "skills": "Навыки"},
        "en": {"profile": "Profile", "experience": "Experience",
               "education": "Education", "skills": "Skills"},
    }
    labels = section_labels.get(lang, section_labels["ru"])

    # --- ПРОФИЛЬ ---
    summary = resume_data.get("summary", "").strip()
    if summary:
        _add_section_header(doc, labels["profile"])
        _add_simple_paragraph(doc, summary, space_after=10)

    # --- ОПЫТ ---
    experience = resume_data.get("experience") or []
    if experience:
        _add_section_header(doc, labels["experience"])
        for role in experience:
            role_title = role.get("role_title", "").strip()
            company_line = role.get("company_line", "").strip()
            role_summary = role.get("role_summary", "").strip()
            bullets = role.get("bullets") or []

            if role_title:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(9)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(role_title)
                _set_run_font(run, 11, bold=True, color=COLOR_HEADING)

            if company_line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(company_line)
                _set_run_font(run, 10, italic=True, color=COLOR_MUTED)

            if role_summary:
                _add_simple_paragraph(doc, role_summary, space_after=5)

            for bullet in bullets:
                if isinstance(bullet, dict):
                    bold_part = bullet.get("bold", "")
                    text = bullet.get("text", "")
                    _add_bullet(doc, bold_part, text)
                elif isinstance(bullet, str):
                    _add_bullet(doc, "", bullet)

    # --- ОБРАЗОВАНИЕ ---
    education = resume_data.get("education") or []
    if education:
        _add_section_header(doc, labels["education"])
        for edu in education:
            degree = edu.get("degree", "").strip() if isinstance(edu, dict) else str(edu)
            institution = edu.get("institution", "").strip() if isinstance(edu, dict) else ""
            if degree:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(degree)
                _set_run_font(run, 10, bold=True, color=COLOR_HEADING)
            if institution:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(institution)
                _set_run_font(run, 10, italic=True, color=COLOR_MUTED)

    # --- НАВЫКИ ---
    skills = resume_data.get("skills") or []
    if skills:
        _add_section_header(doc, labels["skills"])
        for skill in skills:
            if isinstance(skill, dict):
                bold_part = skill.get("bold", "")
                text = skill.get("text", "")
                _add_bullet(doc, bold_part, text)
            elif isinstance(skill, str):
                _add_bullet(doc, "", skill)

    # Возвращаем bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class ResumeService:
    def __init__(self, claude: ClaudeService | None = None):
        self.claude = claude or ClaudeService(model=RESUME_MODEL)

    async def generate(self, profile_data: dict, vacancy: Vacancy) -> ResumeResult:
        user_message = _build_user_message(vacancy, profile_data)

        system_prompt = RESUME_SYSTEM_PROMPT.format(
            current_date=datetime.now(timezone.utc).strftime("%d.%m.%Y")
        )
        response = await self.claude.chat(
            messages=[{"role": "user", "content": user_message}],
            system=system_prompt,
            max_tokens=4096,
            model=RESUME_MODEL,
        )

        raw = (response.text or "").strip()
        if not raw:
            raise ValueError("Claude returned empty resume response")

        # Убираем возможные markdown-фенсы вокруг JSON
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

        try:
            resume_data = json.loads(raw)
        except json.JSONDecodeError as e:
            log.error("resume_json_parse_failed", raw=raw[:500], error=str(e))
            raise ValueError(f"Claude returned invalid JSON: {e}") from e

        docx_bytes = _build_docx(resume_data)

        diagnostics = resume_data.get("diagnostics") or {}
        return ResumeResult(
            docx_bytes=docx_bytes,
            match_percent=int(diagnostics.get("match_percent", 0) or 0),
            strong_alignment=list(diagnostics.get("strong_alignment") or []),
            gaps=list(diagnostics.get("gaps") or []),
            language=resume_data.get("language", "ru"),
        )